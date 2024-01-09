from docx import Document
import pandas as pd


def word_to_excel(word_path, excel_path):
    word_lis = []
    doc = Document(word_path)
    for tb in doc.tables:  # 遍历多个表
        for row in tb.rows:  # 遍历表的每一行
            li = []
            for cell in row.cells:  # 遍历每一行的单元格
                text = ""
                for p in cell.paragraphs:  # 遍历单元格中的每一段（回车符）
                    text += p.text
                li.append(text)
            word_lis.append(li)

    # list转dataframe
    df = pd.DataFrame(word_lis)

    # 保存到本地excel
    df.to_excel(excel_path, index=False, header=0)  # 若不需要将第一行设为表头 header = None


if __name__ == '__main__':
    # 执行方法
    word_path = './zjc.docx'
    excel_path = './zzjc.xlsx'

    word_to_excel(word_path, excel_path)
